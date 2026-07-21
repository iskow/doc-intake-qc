#!/usr/bin/env python3
"""Generate the mock "client dump" fixture for the Document Intake QC Agent.

Why a generator script instead of hand-made files:
- The fixture and the seeded-errors answer key are built from the SAME data
  in this file, so they cannot drift apart.
- Anyone can delete mock-intake/ and rerun this to rebuild it — including
  after a fresh git clone (git does not preserve file modification times,
  so the two date-anomaly files need this script to restore their fake dates).

Outputs (relative to the project root):
- mock-intake/            the messy client dump (~41 files, nested folders)
- seeded-errors.md        human-readable answer key
- seeded-errors.json      machine-readable answer key (used by the QC gate)

Run:  python3 scripts/make_mock_data.py
"""

from __future__ import annotations

import json
import os
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from PIL import Image
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

# --- paths -----------------------------------------------------------------
# pathlib teaching note: Path(__file__) is this script's own location;
# .resolve() makes it absolute; .parents[1] climbs two levels
# (scripts/make_mock_data.py -> scripts -> project root).
ROOT = Path(__file__).resolve().parents[1]
INTAKE = ROOT / "mock-intake"

# --- document authors (embedded metadata; a separate axis from custodian) --
# The "author" is who created a file, stored in the document's own properties
# (PDF /Author, DOCX/XLSX core.xml creator). It is DELIBERATELY distinct from:
#   - the party named in the filename (e.g. "AcmeSupply"), and
#   - the custodian (the collection source — a Phase 4 mapping concept).
# One file below is a planted divergence: an Acme-named contract authored by
# Meridian staff, so the demo can PROVE author != party != custodian. Files
# whose format can't carry an author (txt/csv/png/zip/legacy .doc), scanned
# PDFs, and empty files are left unassigned on purpose.
AUTH_ACME = "Jorge Reyes"       # Acme Supply Co. contact
AUTH_BRIGHT = "Alia Tan"        # BrightWorks Consulting contact
AUTH_DELTA = "Sam Ocampo"       # Delta Freight Ltd. contact
AUTH_INTAKE = "Dana Cruz"       # Meridian Legal — intake / contracts
AUTH_OPS = "M. Villanueva"      # Meridian Legal — operations / reports

# --- document classes (Phase 3 ground truth) -------------------------------
# What each fixture file ACTUALLY is, judged from the content written below --
# not from its folder and not from its name. Phase 3's classifier is scored
# against this map, so the two traps matter: `Contracts/scan_001.pdf` is a
# contract wearing a scanner-default name, and `report_final.pdf` is a report
# sitting at the root as plain text with a .pdf extension.
#
# Some files are honestly ambiguous, so the value is a LIST of acceptable
# labels. `notes.txt` is a client handover note -- a reviewer could file it as
# correspondence or as other, and both are defensible. Scoring the model
# against one arbitrary pick would be fake precision, so we accept either.
#
# None means "no readable text, must NOT be given a label" -- images, the
# archive, the empty files, and the binary junk file. Phase 3 must report
# those as unclassified rather than guess.
#
# Keyed by path relative to mock-intake/. main() asserts this covers every
# generated file, so adding a file without classifying it fails loudly here
# instead of silently skewing the Phase 3 accuracy score.
DOC_CLASSES: dict[str, list[str] | None] = {
    # Invoices -- vendor bills. The spreadsheet is invoice DATA rather than an
    # invoice itself, so "other" is also acceptable for it.
    "Invoices/INV-2026-001_AcmeSupply.pdf": ["invoice"],
    "Invoices/INV-2026-002_AcmeSupply.pdf": ["invoice"],
    "Invoices/INV-2026-003_BrightWorks.pdf": ["invoice"],
    "Invoices/inv 2026-004 brightworks FINAL.pdf": ["invoice"],
    "Invoices/INV-2026-005_AcmeSupply.pdf": ["invoice"],
    "Invoices/INV-2026-005_AcmeSupply (1).pdf": ["invoice"],
    "Invoices/INV-2026-006_DeltaFreight.PDF": ["invoice"],
    "Invoices/INV-2026-007_DeltaFreight.pdf": ["invoice"],
    "Invoices/INV-2026-008_BrightWorks.pdf": ["invoice"],
    "Invoices/INV-2026-009_AcmeSupply.pdf": ["invoice"],
    "Invoices/invoice_batch_Q2.xlsx": ["invoice", "other"],
    # Contracts -- agreements, including the amendment hiding under a scan name.
    "Contracts/CTR-2026-01_ServiceAgreement_AcmeSupply.docx": ["contract"],
    "Contracts/CTR-2026-02_NDA_BrightWorks.docx": ["contract"],
    "Contracts/CTR-2026-02_NDA_BrightWorks_v2.docx": ["contract"],
    "Contracts/CTR-2026-03_MSA_DeltaFreight.pdf": ["contract"],
    "Contracts/Contract FINAL FINAL (use this one).docx": ["contract"],
    "Contracts/scan_001.pdf": ["contract"],                     # TRAP: name says scan
    # Correspondence -- emails rendered to PDF. The empty one has no content.
    "Correspondence/2026-03-14_AcmeSupply_kickoff.pdf": ["correspondence"],
    "Correspondence/2026-03-15_AcmeSupply_followup.pdf": ["correspondence"],
    "Correspondence/2026-04-02_BrightWorks_scope_change.pdf": ["correspondence"],
    "Correspondence/RE RE FW Important!!.pdf": ["correspondence"],
    "Correspondence/memo_draft.docx": None,                     # 0 bytes
    # Reports -- period summaries, including the financial workbook.
    "Reports/RPT-2026-Q1_Operations.pdf": ["report"],
    "Reports/Copy of RPT-2026-Q1_Operations.pdf": ["report"],
    "Reports/RPT-2026-Q2_Operations_DRAFT.docx": ["report"],
    "Reports/RPT-2026-Q1_Financials.xlsx": ["report"],
    # Site photos -- solid-color images, no text of any kind.
    "Site Photos/site_photo_01.png": None,
    "Site Photos/site_photo_02.png": None,
    "Site Photos/site_photo_02 - Copy.png": None,
    "Site Photos/site_photo_03.png": None,
    "Site Photos/IMG_3847.jpg": None,
    # Loose files at the intake root.
    "notes.txt": ["correspondence", "other"],                   # ambiguous on purpose
    "data_export.csv": ["other"],                               # a vendor contact table
    "report_final.pdf": ["report"],                             # TRAP: text, not a PDF
    "empty_placeholder.pdf": None,                              # 0 bytes
    "~$ntract_temp.docx": ["other"],                            # lock-file junk text
    "Thumbs.db": None,                                          # binary junk
    "old_backup.zip": None,                                     # never expanded
    # Deeply nested legacy files.
    "Old Files/2019 archive/deep/nested/legacy_notes.doc": ["other"],
    "Old Files/2019 archive/deep/nested/misc_scan_0042.pdf": ["other"],
    "Old Files/2019 archive/deep/nested/RPT-2031_forecast.pdf": ["report"],
}

# --- the answer key, built as we go ----------------------------------------
seeded: list[dict] = []

# path (relative, posix) -> embedded author. Populated by the builders at the
# moment each file is written, so the key can't drift from what's embedded.
authors: dict[str, str] = {}


def _record_author(path: Path, author: str) -> None:
    """Note a file's embedded author for the answer key, keyed by the same
    relative posix path that seed() and the manifest use."""
    authors[path.relative_to(ROOT).as_posix()] = author


def seed(error_type: str, note: str, *paths: Path) -> None:
    """Record a planted problem. Called at the moment the problem is created,
    so the answer key is a byproduct of building the fixture — not a
    separately maintained list that could go stale."""
    seeded.append(
        {
            "type": error_type,
            "paths": [p.relative_to(ROOT).as_posix() for p in paths],
            "note": note,
        }
    )


# --- file builders ---------------------------------------------------------
def make_pdf(path: Path, title: str, lines: list[str], author: str | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(path), pagesize=LETTER)
    if author:
        c.setAuthor(author)          # writes /Author into the PDF info dict
        _record_author(path, author)
    else:
        c.setAuthor("")              # reportlab defaults /Author to 'anonymous';
        #                              blank it so unauthored PDFs read back as None
    y = 750
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, y, title)
    c.setFont("Helvetica", 11)
    for line in lines:
        y -= 18
        c.drawString(72, y, line)
    c.save()


def make_docx(path: Path, title: str, paragraphs: list[str], author: str | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if author:
        doc.core_properties.author = author   # embedded author (docProps/core.xml)
        _record_author(path, author)
    doc.add_heading(title, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))


def make_xlsx(path: Path, headers: list[str], rows: list[list], author: str | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    if author:
        wb.properties.creator = author        # embedded author (docProps/core.xml)
        _record_author(path, author)
    ws = wb.active
    ws.append(headers)
    for r in rows:
        ws.append(r)
    wb.save(str(path))


def make_png(path: Path, color: tuple[int, int, int]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    Image.new("RGB", (320, 240), color).save(str(path), format="PNG")


def make_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def dupe(src: Path, dst: Path) -> None:
    """Exact duplicate: copy2 preserves content bytes AND the modified time,
    so the pair is hash-identical — exactly what real copy-paste dupes look like.
    The embedded author lives inside those bytes, so the copy inherits it too;
    we mirror that in the answer key."""
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)
    src_rel = src.relative_to(ROOT).as_posix()
    if src_rel in authors:
        _record_author(dst, authors[src_rel])


def set_mtime(path: Path, when: datetime) -> None:
    """os.utime sets (access_time, modified_time). Creation time is not
    settable from Python on most systems, so date anomalies are seeded via
    the MODIFIED date only — Phase 2 rules should check mtime."""
    ts = when.timestamp()
    os.utime(path, (ts, ts))


# --- invoice text helper (gives Phase 3 something real to classify) --------
def invoice_lines(no: str, vendor: str, amount: str, date: str) -> list[str]:
    return [
        f"Invoice No: {no}",
        f"Vendor: {vendor}",
        "Bill To: Meridian Legal Services LLC",
        f"Invoice Date: {date}",
        f"Amount Due: {amount}",
        "Payment Terms: Net 30",
    ]


def main() -> None:
    if INTAKE.exists():
        shutil.rmtree(INTAKE)

    inv = INTAKE / "Invoices"
    ctr = INTAKE / "Contracts"
    cor = INTAKE / "Correspondence"
    rpt = INTAKE / "Reports"
    pho = INTAKE / "Site Photos"

    # ------------------------------------------------------------ Invoices
    # Vendor invoices carry the vendor contact as author; the internal batch
    # spreadsheet is authored by Meridian (author != any single party).
    make_pdf(inv / "INV-2026-001_AcmeSupply.pdf", "INVOICE",
             invoice_lines("INV-2026-001", "Acme Supply Co.", "$1,250.00", "2026-03-02"),
             author=AUTH_ACME)
    make_pdf(inv / "INV-2026-002_AcmeSupply.pdf", "INVOICE",
             invoice_lines("INV-2026-002", "Acme Supply Co.", "$980.50", "2026-03-19"),
             author=AUTH_ACME)
    make_pdf(inv / "INV-2026-003_BrightWorks.pdf", "INVOICE",
             invoice_lines("INV-2026-003", "BrightWorks Consulting", "$4,200.00", "2026-03-28"),
             author=AUTH_BRIGHT)

    p = inv / "inv 2026-004 brightworks FINAL.pdf"
    make_pdf(p, "INVOICE",
             invoice_lines("INV-2026-004", "BrightWorks Consulting", "$1,875.00", "2026-04-06"),
             author=AUTH_BRIGHT)
    seed("naming-violation", "Spaces, lowercase, and 'FINAL' suffix — breaks INV-YYYY-NNN_Vendor convention.", p)

    src = inv / "INV-2026-005_AcmeSupply.pdf"
    make_pdf(src, "INVOICE",
             invoice_lines("INV-2026-005", "Acme Supply Co.", "$2,310.00", "2026-04-15"),
             author=AUTH_ACME)
    d = inv / "INV-2026-005_AcmeSupply (1).pdf"
    dupe(src, d)
    seed("exact-duplicate", "Copy-paste artifact — identical bytes, ' (1)' name.", src, d)

    p = inv / "INV-2026-006_DeltaFreight.PDF"
    make_pdf(p, "INVOICE",
             invoice_lines("INV-2026-006", "Delta Freight Ltd.", "$640.00", "2026-04-22"),
             author=AUTH_DELTA)
    seed("naming-violation", "Uppercase '.PDF' extension — inconsistent with the rest of the set.", p)

    make_pdf(inv / "INV-2026-007_DeltaFreight.pdf", "INVOICE",
             invoice_lines("INV-2026-007", "Delta Freight Ltd.", "$1,120.00", "2026-05-01"),
             author=AUTH_DELTA)
    make_pdf(inv / "INV-2026-008_BrightWorks.pdf", "INVOICE",
             invoice_lines("INV-2026-008", "BrightWorks Consulting", "$3,050.00", "2026-05-09"),
             author=AUTH_BRIGHT)
    make_pdf(inv / "INV-2026-009_AcmeSupply.pdf", "INVOICE",
             invoice_lines("INV-2026-009", "Acme Supply Co.", "$775.25", "2026-05-20"),
             author=AUTH_ACME)
    make_xlsx(inv / "invoice_batch_Q2.xlsx",
              ["invoice_no", "vendor", "date", "amount_usd"],
              [["INV-2026-004", "BrightWorks Consulting", "2026-04-06", 1875.00],
               ["INV-2026-005", "Acme Supply Co.", "2026-04-15", 2310.00],
               ["INV-2026-006", "Delta Freight Ltd.", "2026-04-22", 640.00],
               ["INV-2026-007", "Delta Freight Ltd.", "2026-05-01", 1120.00]],
              author=AUTH_INTAKE)

    # ------------------------------------------------------------ Contracts
    # DIVERGENCE FILE: filename party = Acme, but Meridian drafted it, so the
    # embedded author is Dana Cruz (Meridian). This one file proves the demo's
    # whole point — author != party != custodian.
    make_docx(ctr / "CTR-2026-01_ServiceAgreement_AcmeSupply.docx", "SERVICE AGREEMENT",
              ["This Service Agreement (the 'Agreement') is entered into as of March 1, 2026, "
               "by and between Meridian Legal Services LLC and Acme Supply Co.",
               "1. Services. The Vendor shall provide office supply and logistics services.",
               "2. Term. This Agreement runs for twelve (12) months from the effective date.",
               "3. Fees. Fees are payable net 30 from receipt of a valid invoice."],
              author=AUTH_INTAKE)
    src = ctr / "CTR-2026-02_NDA_BrightWorks.docx"
    make_docx(src, "MUTUAL NON-DISCLOSURE AGREEMENT",
              ["This Mutual Non-Disclosure Agreement is made between Meridian Legal Services LLC "
               "and BrightWorks Consulting, effective March 10, 2026.",
               "Each party agrees to hold the other's Confidential Information in strict confidence.",
               "Term: three (3) years from the effective date."],
              author=AUTH_INTAKE)
    p = ctr / "CTR-2026-02_NDA_BrightWorks_v2.docx"
    make_docx(p, "MUTUAL NON-DISCLOSURE AGREEMENT (REV. 2)",
              ["This Mutual Non-Disclosure Agreement is made between Meridian Legal Services LLC "
               "and BrightWorks Consulting, effective March 10, 2026.",
               "Each party agrees to hold the other's Confidential Information in strict confidence.",
               "Term: five (5) years from the effective date. Adds clause 4 (residuals)."],
              author=AUTH_INTAKE)
    seed("near-duplicate-name", "'_v2' variant beside the original — same base name, different content. "
         "Which version is authoritative?", src, p)

    p = ctr / "Contract FINAL FINAL (use this one).docx"
    make_docx(p, "CONSULTING CONTRACT",
              ["Consulting contract between Meridian Legal Services LLC and Delta Freight Ltd.",
               "Scope: quarterly logistics review and reporting.",
               "Fee: $2,000 per quarter, invoiced in arrears."],
              author=AUTH_INTAKE)
    seed("naming-violation", "'FINAL FINAL (use this one)' — no ID, no convention, ambiguous authority.", p)

    make_pdf(ctr / "CTR-2026-03_MSA_DeltaFreight.pdf", "MASTER SERVICE AGREEMENT",
             ["Master Service Agreement between Meridian Legal Services LLC and Delta Freight Ltd.",
              "Effective Date: April 1, 2026.",
              "Governs all statements of work executed under this MSA."],
             author=AUTH_INTAKE)

    p = ctr / "scan_001.pdf"
    make_pdf(p, "AMENDMENT NO. 1",
             ["Amendment No. 1 to the Service Agreement dated March 1, 2026,",
              "between Meridian Legal Services LLC and Acme Supply Co.",
              "Section 3 (Fees) is amended to add a 2% early-payment discount."])
    seed("naming-violation", "'scan_001.pdf' — uninformative scanner-default name; content is a contract amendment.", p)

    # ------------------------------------------------------- Correspondence
    # Emails rendered to PDF: most have no author (exported/printed), one
    # carries the sender — a realistic mix that feeds the Unassigned bucket.
    make_pdf(cor / "2026-03-14_AcmeSupply_kickoff.pdf", "EMAIL",
             ["From: j.reyes@acmesupply.example", "To: intake@meridianlegal.example",
              "Date: 2026-03-14", "Subject: Kickoff — supply agreement onboarding",
              "Hi team, attaching the signed agreement and our W-9. Looking forward to starting Monday."])
    make_pdf(cor / "2026-03-15_AcmeSupply_followup.pdf", "EMAIL",
             ["From: intake@meridianlegal.example", "To: j.reyes@acmesupply.example",
              "Date: 2026-03-15", "Subject: RE: Kickoff — supply agreement onboarding",
              "Thanks Jorge — received. We'll confirm the PO number by Wednesday."])
    make_pdf(cor / "2026-04-02_BrightWorks_scope_change.pdf", "EMAIL",
             ["From: a.tan@brightworks.example", "To: intake@meridianlegal.example",
              "Date: 2026-04-02", "Subject: Scope change request — Q2 engagement",
              "Requesting an additional workshop in May; revised SOW to follow."],
             author=AUTH_BRIGHT)

    p = cor / "RE RE FW Important!!.pdf"
    make_pdf(p, "EMAIL",
             ["From: unknown@deltafreight.example", "To: intake@meridianlegal.example",
              "Date: 2026-04-11", "Subject: RE: RE: FW: Important!!",
              "Please see the thread below regarding the delayed shipment."])
    seed("naming-violation", "Forward-chain subject as filename — spaces, '!!', no date, no sender.", p)

    p = cor / "memo_draft.docx"
    p.parent.mkdir(parents=True, exist_ok=True)
    p.touch()  # touch() creates an empty file — 0 bytes
    seed("zero-byte", "Empty .docx — likely a failed save or interrupted transfer.", p)

    # -------------------------------------------------------------- Reports
    src = rpt / "RPT-2026-Q1_Operations.pdf"
    make_pdf(src, "Q1 2026 OPERATIONS REPORT",
             ["Prepared by: Meridian Legal Services LLC — Operations",
              "Period: January-March 2026",
              "Intake volume up 12% quarter over quarter.",
              "Average processing turnaround: 2.4 business days."],
             author=AUTH_OPS)
    d = rpt / "Copy of RPT-2026-Q1_Operations.pdf"
    dupe(src, d)
    seed("exact-duplicate", "'Copy of' artifact — identical bytes under a different name.", src, d)

    make_docx(rpt / "RPT-2026-Q2_Operations_DRAFT.docx", "Q2 2026 OPERATIONS REPORT (DRAFT)",
              ["Period: April-June 2026 (draft — figures not final).",
               "Intake volume tracking flat; two vendor onboardings completed."],
              author=AUTH_OPS)
    make_xlsx(rpt / "RPT-2026-Q1_Financials.xlsx",
              ["month", "revenue_usd", "expenses_usd"],
              [["2026-01", 42000, 31000], ["2026-02", 45500, 32200], ["2026-03", 47800, 33900]],
              author=AUTH_OPS)

    # ---------------------------------------------------------- Site Photos
    make_png(pho / "site_photo_01.png", (70, 130, 180))
    src = pho / "site_photo_02.png"
    make_png(src, (34, 139, 34))
    d = pho / "site_photo_02 - Copy.png"
    dupe(src, d)
    seed("exact-duplicate", "' - Copy' artifact — identical bytes.", src, d)
    make_png(pho / "site_photo_03.png", (178, 34, 34))

    p = pho / "IMG_3847.jpg"
    make_png(p, (218, 165, 32))  # PNG bytes saved under a .jpg name
    seed("extension-mismatch", "Named .jpg but the content is PNG (magic bytes say PNG).", p)

    # ---------------------------------------------------------- Root, loose
    make_text(INTAKE / "notes.txt",
              "Handover notes from client: invoices for Q1-Q2, contracts folder may contain older versions, "
              "photos are from the April site visit. Backup zip includes files from the old drive.\n")
    make_text(INTAKE / "data_export.csv",
              "vendor,contact,email\nAcme Supply Co.,Jorge Reyes,j.reyes@acmesupply.example\n"
              "BrightWorks Consulting,Alia Tan,a.tan@brightworks.example\n"
              "Delta Freight Ltd.,Sam Ocampo,s.ocampo@deltafreight.example\n")

    p = INTAKE / "report_final.pdf"
    make_text(p, "Q4 2025 summary (plain text pasted into a file and renamed):\n"
                 "Total intake 1,204 documents; 3 vendors active; no open exceptions at year end.\n")
    seed("extension-mismatch", "Named .pdf but the content is plain text — no PDF header.", p)
    seed("naming-violation", "'report_final.pdf' — uninformative name, no ID or date.", p)

    p = INTAKE / "empty_placeholder.pdf"
    p.touch()
    seed("zero-byte", "Empty .pdf — 0 bytes.", p)

    p = INTAKE / "~$ntract_temp.docx"
    make_text(p, "Office lock-file junk content")
    seed("junk-file", "Office temp/lock file ('~$' prefix) — should never ship in a client dump.", p)

    p = INTAKE / "Thumbs.db"
    p.write_bytes(b"\xd0\xcf\x11\xe0dummy-thumbs-db")
    seed("junk-file", "Windows thumbnail-cache system file.", p)

    zp = INTAKE / "old_backup.zip"
    with zipfile.ZipFile(zp, "w") as z:
        z.writestr("old_drive/vendor_list_2025.txt",
                   "Acme Supply Co.\nBrightWorks Consulting\n")
        z.writestr("old_drive/meeting_notes_2025-11.txt",
                   "Notes from the November 2025 vendor review meeting.\n")
    seed("archive", "ZIP present in the intake — contents (2 files) are invisible until expanded.", zp)

    # -------------------------------------------------- Deeply nested + dates
    deep = INTAKE / "Old Files" / "2019 archive" / "deep" / "nested"
    p = deep / "misc_scan_0042.pdf"
    make_pdf(p, "SCANNED DOCUMENT",
             ["Legacy scanned page recovered from the 2019 archive drive.",
              "Content: vendor onboarding checklist (superseded)."])
    seed("deep-nesting", "Four folder levels down — easy to miss in a manual review.", p)

    p = deep / "legacy_notes.doc"
    make_text(p, "Legacy notes migrated from the old system. Plain text saved with a .doc name.\n")
    set_mtime(p, datetime(1980, 5, 12, 9, 30))
    seed("extension-mismatch", "Named .doc but the content is plain text.", p)
    seed("date-anomaly", "Modified date 1980-05-12 — predates the client relationship (and most PCs).", p)

    p = deep / "RPT-2031_forecast.pdf"
    make_pdf(p, "FORECAST 2031",
             ["Long-range forecast document.", "Placeholder projections for 2031."],
             author=AUTH_OPS)
    set_mtime(p, datetime(2031, 1, 15, 12, 0))
    seed("date-anomaly", "Modified date 2031-01-15 — in the future; clock error or metadata tampering.", p)

    # ------------------------------------------------------- Write answer key
    all_files = sorted(q.relative_to(ROOT).as_posix() for q in INTAKE.rglob("*") if q.is_file())
    # Full path -> author map, including None for files with no embedded author
    # (Unassigned). Ground truth for Phase 2's author-extraction check.
    authors_map = {f: authors.get(f) for f in all_files}

    # Document-class ground truth, keyed on the same full relative path as
    # everything else. The two set-difference checks below are the drift guard:
    # a new fixture file with no class, or a class entry naming a file that no
    # longer exists, stops the generator instead of quietly skewing Phase 3's
    # accuracy score.
    classes_map = {f: DOC_CLASSES.get(f[len("mock-intake/"):]) for f in all_files}
    unclassified = sorted(f for f in all_files
                          if f[len("mock-intake/"):] not in DOC_CLASSES)
    stale = sorted(k for k in DOC_CLASSES if f"mock-intake/{k}" not in all_files)
    if unclassified or stale:
        raise SystemExit(
            "DOC_CLASSES is out of sync with the fixture.\n"
            + "".join(f"  missing a class: {f}\n" for f in unclassified)
            + "".join(f"  names a file that does not exist: {k}\n" for k in stale)
        )

    key = {
        "generated": datetime.now().isoformat(timespec="seconds"),
        "file_count": len(all_files),
        "files": all_files,
        "seeded_errors": seeded,
        "authors": authors_map,
        "classes": classes_map,
    }
    (ROOT / "seeded-errors.json").write_text(json.dumps(key, indent=2), encoding="utf-8")

    by_type: dict[str, list[dict]] = {}
    for s in seeded:
        by_type.setdefault(s["type"], []).append(s)

    lines = [
        "# seeded-errors.md — Answer key for the mock intake fixture",
        "",
        "_Generated by `scripts/make_mock_data.py` — do not edit by hand; rerun the script instead._",
        f"_Generated: {key['generated']} · Files in fixture: {key['file_count']} "
        "(+2 inside old_backup.zip)_",
        "",
        "**Notes:**",
        "- Git does not preserve modified times. After a fresh clone, rerun the generator "
        "to restore the two date anomalies.",
        "- PDF bytes embed a creation timestamp, so hashes differ between generator runs — "
        "but within any one run, each duplicate pair is hash-identical.",
        "- Date anomalies are seeded on the MODIFIED date (creation time is not settable from Python).",
        "",
    ]
    for etype in sorted(by_type):
        entries = by_type[etype]
        lines.append(f"## {etype} ({len(entries)})")
        lines.append("")
        for e in entries:
            for path in e["paths"]:
                lines.append(f"- `{path}`")
            lines.append(f"  - {e['note']}")
        lines.append("")

    # Document authors are metadata, not errors — their own section.
    named = {f: a for f, a in authors_map.items() if a}
    lines.append(f"## document-authors ({len(named)} of {len(all_files)} files)")
    lines.append("")
    lines.append("_Embedded author metadata — a SEPARATE axis from custodian (collection "
                 "source) and from the party in the filename. Full path->author map "
                 "(including nulls) is in seeded-errors.json._")
    lines.append("")
    for f in sorted(named):
        lines.append(f"- `{f}` -> {named[f]}")
    lines.append("")
    lines.append(f"_{len(all_files) - len(named)} files are Unassigned — their format "
                 "can't carry an author (txt/csv/png/zip/legacy .doc), they're scanned, "
                 "or they're empty._")
    lines.append("- **Divergence check:** `mock-intake/Contracts/CTR-2026-01_ServiceAgreement_"
                 "AcmeSupply.docx` — filename party = Acme, embedded author = Dana Cruz "
                 "(Meridian). Proof that author != party != custodian.")
    lines.append("")

    # Document classes are ground truth for Phase 3, not errors — own section.
    classifiable = {f: c for f, c in classes_map.items() if c}
    lines.append(f"## document-classes ({len(classifiable)} of {len(all_files)} files classifiable)")
    lines.append("")
    lines.append("_What each file ACTUALLY is, judged from its content — independent of "
                 "the folder it sits in and the name it wears. Phase 3's classifier is "
                 "scored against this. Files listed with two labels are honestly "
                 "ambiguous and either answer is accepted._")
    lines.append("")
    for f in sorted(classifiable):
        lines.append(f"- `{f}` -> {' | '.join(classifiable[f])}")
    lines.append("")
    lines.append(f"_{len(all_files) - len(classifiable)} files carry no readable text "
                 "(images, the archive, the empty files, binary junk). Phase 3 must "
                 "report these as **unclassified** — never guess a label._")
    lines.append("")
    lines.append("- **Class traps:** `mock-intake/Contracts/scan_001.pdf` is a contract "
                 "amendment under a scanner-default name, and `mock-intake/report_final.pdf` "
                 "is a report sitting at the root as plain text with a `.pdf` extension. "
                 "Folder and filename are evidence, not gospel.")
    lines.append("")
    (ROOT / "seeded-errors.md").write_text("\n".join(lines), encoding="utf-8")

    print(f"Fixture built: {key['file_count']} files in {INTAKE.name}/ (+2 inside the zip)")
    print(f"Seeded errors: {len(seeded)} across {len(by_type)} types: "
          + ", ".join(f"{t}={len(v)}" for t, v in sorted(by_type.items())))
    print(f"Authors embedded: {len(named)} files "
          f"({len(all_files) - len(named)} unassigned).")
    print(f"Classes assigned: {len(classifiable)} classifiable files "
          f"({len(all_files) - len(classifiable)} with no readable text).")


if __name__ == "__main__":
    main()
